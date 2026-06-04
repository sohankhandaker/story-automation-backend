import io
import logging
import re
from pathlib import Path
from datetime import datetime
from fastapi import APIRouter, Depends, BackgroundTasks, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified
from ..database import get_db
from .. import models, schemas
from ..deps import get_current_user
from ..services import agent
from ..services.crawler import fetch_page_text
from ..services.github import cfg_for_user, cfg_for_project
from ..services import github as gh

log = logging.getLogger(__name__)
router = APIRouter(prefix="/api/notes", tags=["notes"])
projects_notes_router = APIRouter(prefix="/api/projects", tags=["notes"])


# ── Helpers ───────────────────────────────────────────────────────────────────

def _title_from_raw(raw: str) -> str:
    """Extract a readable title from the first non-empty line of raw notes (max 80 chars)."""
    for line in raw.splitlines():
        line = line.strip()
        if line:
            return line[:80] + ('…' if len(line) > 80 else '')
    return "New Note"


def _note_dict(note: models.MeetingNote, db: Session) -> dict:
    """Return a note serialised to dict, enriched with PRD info from the linked PrdDocument."""
    prd = db.query(models.PrdDocument).filter(models.PrdDocument.note_id == note.id).first()
    return {
        "id": note.id,
        "project_id": note.project_id,
        "note_type": note.note_type or "note",
        "title": note.title,
        "raw_notes": note.raw_notes,
        "wiki_url": note.wiki_url,
        "brd_draft": note.brd_draft,
        "brd_generation_phase": note.brd_generation_phase,
        "status": note.status,
        "current_version_number": note.current_version_number,
        "github_issue_url": note.github_issue_url,
        "github_issue_number": note.github_issue_number,
        "reviewer_github_username": note.reviewer_github_username,
        "reviewer_name": note.reviewer_name,
        "reviewers": note.reviewers or [],
        "github_file_url": note.github_file_url,
        "github_file_raw_url": note.github_file_raw_url,
        "planner_doc_content": note.planner_doc_content,
        "planner_doc_url": note.planner_doc_url,
        "prd_status": prd.status if prd else None,
        "prd_version_number": prd.current_version_number if prd else None,
        "prd_file_url": prd.github_file_url if prd else None,
        "prd_file_raw_url": prd.github_file_raw_url if prd else None,
        "created_at": note.created_at,
        "updated_at": note.updated_at,
    }


def _combined_notes(note: models.MeetingNote) -> str:
    """Combine all note entries and attachment text chronologically."""
    entries = sorted(note.entries or [], key=lambda e: e.created_at)
    parts = []
    if entries:
        parts.extend(
            f"[Entry {i + 1} — {e.created_at.strftime('%Y-%m-%d %H:%M')}]\n{e.content}"
            for i, e in enumerate(entries)
        )
    else:
        parts.append(note.raw_notes)
    for att in (note.attachments or []):
        if att.content_text:
            parts.append(f"[Attachment: {att.filename}]\n{att.content_text}")
    return "\n\n---\n\n".join(parts)


def _get_project_issue_number(note: models.MeetingNote, db: Session) -> int | None:
    """Return the GitHub issue number from the note's project, if any."""
    if not note.project_id:
        return note.github_issue_number
    project = db.query(models.Project).filter(models.Project.id == note.project_id).first()
    return project.github_issue_number if project else None


def _build_brd_comment(note: models.MeetingNote, brd_markdown: str) -> str:
    title = note.title or "BRD"
    header = (
        f"## BRD: {title} — v{note.current_version_number}\n\n"
        f"---\n\n"
    )
    body = header + brd_markdown
    if len(body) > 60000:
        body = body[:60000] + "\n\n_(Content truncated — see app for full BRD)_"
    if note.github_file_url:
        body += f"\n\n---\n\n📥 **[View / Download BRD (.md)]({note.github_file_url})**"
    return body


# ── Background tasks ──────────────────────────────────────────────────────────

def _full_brd_pipeline(note_id: str):
    """
    Full pipeline triggered by Mark as Ready:
      Phase 1 — crawl all URLs in notes + wiki_url
      Phase 2 — AI analysis of all notes
      Phases 3-9 — 7-phase BRD generation using analysis as context
    Posts progress and result as comments on the project's GitHub issue.
    """
    from ..database import SessionLocal
    db = SessionLocal()
    note = None
    try:
        note = db.query(models.MeetingNote).filter(models.MeetingNote.id == note_id).first()
        if not note:
            return

        creator = db.query(models.User).filter(models.User.id == note.creator_id).first()
        cfg = cfg_for_user(creator)

        issue_number = _get_project_issue_number(note, db)
        combined = _combined_notes(note)

        # Build project + customer context block
        project_context = ""
        if note.project_id:
            project = db.query(models.Project).filter(
                models.Project.id == note.project_id
            ).first()
            if project:
                project_context += (
                    f"\n\n## Project Context\n"
                    f"**Project Name:** {project.title}\n"
                )
                if project.short_description:
                    project_context += f"**Project Description:** {project.short_description}\n"
                if project.url:
                    project_context += f"**Project URL:** {project.url}\n"

                if project.customer_id:
                    customer = db.query(models.Customer).filter(
                        models.Customer.id == project.customer_id
                    ).first()
                    if customer:
                        project_context += (
                            f"\n## Client / Customer Context\n"
                            f"**Client Name:** {customer.name}\n"
                        )
                        if customer.short_description:
                            project_context += f"**Client Description:** {customer.short_description}\n"
                        if customer.url:
                            project_context += f"**Client Website:** {customer.url}\n"

        # Phase 1: Crawl URLs
        note.brd_generation_phase = 1
        db.commit()

        urls_to_crawl = []
        if note.wiki_url:
            for u in note.wiki_url.split('\n'):
                u = u.strip()
                if u:
                    urls_to_crawl.append(u)
        for u in re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', combined)[:3]:
            if u not in urls_to_crawl:
                urls_to_crawl.append(u)

        # Also crawl the customer's website if available
        if note.project_id:
            _proj = db.query(models.Project).filter(
                models.Project.id == note.project_id
            ).first()
            if _proj and _proj.customer_id:
                _cust = db.query(models.Customer).filter(
                    models.Customer.id == _proj.customer_id
                ).first()
                if _cust and _cust.url and _cust.url not in urls_to_crawl:
                    urls_to_crawl.insert(0, _cust.url)  # prioritise customer URL

        crawled_content = ""
        for url in urls_to_crawl[:5]:
            try:
                content = fetch_page_text(url)
                if content:
                    crawled_content += f"\n\n[From {url}]\n{content[:3000]}"
            except Exception as e:
                log.warning(f"Crawl failed for {url}: {e}")

        # Prepend project/customer context to crawled content so the AI has it
        if project_context:
            crawled_content = project_context + crawled_content

        # Phase 2: Analyze notes
        note.brd_generation_phase = 2
        db.commit()

        analysis = agent.analyze_notes_to_draft(combined, crawled_content)

        # Post analysis as a comment so BU can see early progress
        if issue_number:
            try:
                analysis_comment = (
                    f"### BRD In Progress — Notes Analysis\n\n"
                    f"*Generating full BRD — estimated 1–2 minutes…*\n\n"
                    f"---\n\n{analysis}"
                )
                if len(analysis_comment) > 60000:
                    analysis_comment = analysis_comment[:60000] + "\n\n_(truncated)_"
                gh.add_comment(issue_number, analysis_comment, cfg=cfg)
            except Exception as e:
                log.warning(f"GitHub analysis comment failed: {e}")

        # Extract tentative title from analysis
        for line in analysis.splitlines():
            stripped = line.lstrip("#* ").strip("* ").strip()
            if stripped and len(stripped) > 3 and not stripped.lower().startswith("notes analysis"):
                if line.startswith("#"):
                    note.title = stripped
                    db.commit()
                    break

        # Phases 3-9: 7-phase BRD generation (offset +2 from BRD phase number)
        def _phase_cb(phase_num: int, _total: int):
            try:
                note.brd_generation_phase = phase_num + 2
                db.commit()
            except Exception as e:
                log.warning(f"Phase callback commit failed: {e}")

        result = agent.enhance_notes_to_brd(
            combined, crawled_content, analysis=analysis, phase_callback=_phase_cb
        )

        v1 = models.BrdVersion(
            note_id=note.id,
            version_number=1,
            brd_markdown=result["brd_markdown"],
            change_summary="Initial AI generation",
            changed_sections=[],
        )
        db.add(v1)

        note.title = result["title"]
        note.brd_draft = result["brd_markdown"]
        note.brd_generation_phase = None
        note.current_version_number = 1
        note.status = "Pending Review"
        db.commit()

        # Push .md file first so the download link is available in the comment
        safe_title = re.sub(r'[^\w\s\-]', '', note.title or 'BRD').strip().replace(' ', '_')
        file_path = f"brd/{safe_title}_{note.id[:8]}.md"
        try:
            file_info = gh.push_file(
                path=file_path,
                content=result["brd_markdown"],
                commit_message=f"docs(brd): add BRD for {note.title or note.id[:8]}",
                cfg=cfg,
            )
            note.github_file_url = file_info["html_url"]
            note.github_file_raw_url = file_info["raw_url"]
            db.commit()
            log.info(f"BRD .md pushed to repo: {file_info['html_url']}")
        except Exception as e:
            log.warning(f"BRD file push failed: {e}")

        # Post full BRD as a comment (includes download link if file was pushed)
        if issue_number:
            try:
                gh.add_comment(
                    issue_number,
                    _build_brd_comment(note, result["brd_markdown"]),
                    cfg=cfg,
                )
                gh.add_comment(
                    issue_number,
                    "BRD generation complete!\n\nThe full Business Requirements Document is in the comment above.\n\nThe creator will assign a reviewer shortly.",
                    cfg=cfg,
                )
            except Exception as e:
                log.warning(f"GitHub BRD comment failed: {e}")

    except Exception as e:
        log.error(f"BRD pipeline failed for note {note_id}: {e}")
        try:
            if note:
                note.brd_generation_phase = None
                note.status = "Draft"
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


def _full_cr_pipeline(note_id: str):
    """Background: generate CR summary against approved PRD, then set status to In Review."""
    from ..database import SessionLocal
    db = SessionLocal()
    try:
        note = db.query(models.MeetingNote).filter(models.MeetingNote.id == note_id).first()
        if not note or note.note_type != "change_request":
            return

        creator = db.query(models.User).filter(models.User.id == note.creator_id).first()
        # Find the project's approved/sent-to-planner PRD for context
        prd = None
        if note.project_id:
            prd = (
                db.query(models.PrdDocument)
                .join(models.MeetingNote, models.MeetingNote.id == models.PrdDocument.note_id)
                .filter(
                    models.MeetingNote.project_id == note.project_id,
                    models.PrdDocument.status == "Sent to Planner",
                )
                .first()
            )

        prd_content = (prd.prd_draft or "") if prd else ""

        # Gather attachment text
        attachments = db.query(models.NoteAttachment).filter(
            models.NoteAttachment.note_id == note_id
        ).all()
        combined = note.raw_notes
        for att in attachments:
            if att.content_text:
                combined += f"\n\n[Attachment: {att.filename}]\n{att.content_text}"

        summary = agent.generate_cr_summary(combined, prd_content)

        note.brd_draft = summary
        note.brd_generation_phase = None
        note.status = "In Review"
        db.commit()

        # Post to GitHub ticket
        if note.github_issue_number and creator:
            try:
                project = db.query(models.Project).filter(
                    models.Project.id == note.project_id
                ).first() if note.project_id else None
                cfg = cfg_for_project(creator, project) if project else cfg_for_user(creator)
                gh.add_comment(
                    note.github_issue_number,
                    f"### CR Summary Generated\n\n{summary[:1500]}{'…' if len(summary) > 1500 else ''}",
                    cfg=cfg,
                )
            except Exception as e:
                log.warning(f"CR GitHub comment failed: {e}")

        log.info(f"CR pipeline complete for note {note_id}")
    except Exception as e:
        log.error(f"_full_cr_pipeline failed for note {note_id}: {e}")
        try:
            note = db.query(models.MeetingNote).filter(models.MeetingNote.id == note_id).first()
            if note:
                note.brd_generation_phase = None
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


def _cr_send_to_planner_task(note_id: str):
    """Background: generate planner doc, upload to GitHub, set CR status to Closed."""
    from ..database import SessionLocal
    db = SessionLocal()
    try:
        note = db.query(models.MeetingNote).filter(models.MeetingNote.id == note_id).first()
        if not note or note.note_type != "change_request":
            return

        creator = db.query(models.User).filter(models.User.id == note.creator_id).first()
        prd = None
        if note.project_id:
            prd = (
                db.query(models.PrdDocument)
                .join(models.MeetingNote, models.MeetingNote.id == models.PrdDocument.note_id)
                .filter(
                    models.MeetingNote.project_id == note.project_id,
                    models.PrdDocument.status == "Sent to Planner",
                )
                .first()
            )

        prd_content = (prd.prd_draft or "") if prd else ""
        cr_summary = note.brd_draft or note.raw_notes

        planner_doc = agent.generate_cr_planner_doc(
            cr_title=note.title or "Change Request",
            cr_summary=cr_summary,
            prd_content=prd_content,
        )

        note.planner_doc_content = planner_doc
        note.status = "Closed"
        db.commit()

        # Upload planner doc to GitHub repo
        planner_url = None
        if creator:
            project = db.query(models.Project).filter(
                models.Project.id == note.project_id
            ).first() if note.project_id else None
            cfg = cfg_for_project(creator, project) if project else cfg_for_user(creator)
            try:
                safe = note.id.replace("-", "")[:12]
                file_result = gh.push_file(
                    path=f"change-requests/cr-{safe}.md",
                    content=planner_doc,
                    commit_message=f"docs: planner doc for CR {note.title or note.id}",
                    cfg=cfg,
                )
                planner_url = file_result.get("html_url") or file_result.get("raw_url")
                note.planner_doc_url = planner_url
                db.commit()
            except Exception as e:
                log.warning(f"CR planner doc upload failed: {e}")

            # Comment on GitHub ticket with the planner doc link
            if note.github_issue_number:
                try:
                    link_line = f"\n📎 [Planner Document]({planner_url})" if planner_url else ""
                    gh.add_comment(
                        note.github_issue_number,
                        f"### ✅ Change Request Closed — Planner Document Generated{link_line}\n\n"
                        f"The CR has been finalised and a planner handoff document has been created.",
                        cfg=cfg,
                    )
                except Exception as e:
                    log.warning(f"CR closure comment failed: {e}")

        log.info(f"CR send-to-planner complete for note {note_id}")
    except Exception as e:
        log.error(f"_cr_send_to_planner_task failed for {note_id}: {e}")
    finally:
        db.close()


def _update_brd(note_id: str, feedback: str):
    """Update BRD from reviewer/creator feedback, save new version."""
    from ..database import SessionLocal
    db = SessionLocal()
    try:
        note = db.query(models.MeetingNote).filter(models.MeetingNote.id == note_id).first()
        if not note or not note.brd_draft:
            return

        result = agent.update_brd_from_feedback(
            current_brd=note.brd_draft,
            reviewer_comments=[feedback],
        )

        next_ver = (note.current_version_number or 1) + 1
        version = models.BrdVersion(
            note_id=note.id,
            version_number=next_ver,
            brd_markdown=result["updated_markdown"],
            change_summary=result["change_summary"],
            changed_sections=result["changed_sections"],
            reviewer_comment=feedback,
        )
        db.add(version)

        note.brd_draft = result["updated_markdown"]
        note.current_version_number = next_ver
        note.brd_generation_phase = None

        issue_number = _get_project_issue_number(note, db)
        if issue_number:
            try:
                creator = db.query(models.User).filter(
                    models.User.id == note.creator_id
                ).first()
                cfg = cfg_for_user(creator)

                sections_line = ", ".join(result["changed_sections"]) if result["changed_sections"] else "General updates"
                reviewer = f"@{note.reviewer_github_username}" if note.reviewer_github_username else "Reviewer"
                update_comment = (
                    f"## BRD Updated (v{next_ver})\n\n"
                    f"**Changes:** {result['change_summary']}\n\n"
                    f"**Sections updated:** {sections_line}\n\n"
                    f"---\n\n"
                    f"{_build_brd_comment(note, result['updated_markdown'])}\n\n"
                    f"---\n\n"
                    f"{reviewer} — please review the updated BRD above. "
                    f"Reply **`APPROVED`** to approve, or leave further feedback."
                )
                gh.add_comment(issue_number, update_comment, cfg=cfg)
            except Exception as e:
                log.warning(f"GitHub update comment failed after BRD feedback: {e}")

        note.status = "In Review"
        db.commit()

    except Exception as e:
        log.error(f"BRD update failed for note {note_id}: {e}")
        try:
            note.brd_generation_phase = None
            db.commit()
        except Exception:
            pass
    finally:
        db.close()


# ── Project-scoped note routes ─────────────────────────────────────────────────

@projects_notes_router.post("/{project_id}/notes", response_model=schemas.MeetingNoteResponse, status_code=201)
def create_note_for_project(
    project_id: str,
    body: schemas.NotesEnhanceRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    project = db.query(models.Project).filter(
        models.Project.id == project_id,
        models.Project.creator_id == current_user.id,
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    is_cr = (body.note_type or "note") == "change_request"
    note_title = _title_from_raw(body.raw_notes)

    # For change requests create a fresh GitHub issue on the project board
    cr_issue_url = None
    cr_issue_number = None
    cr_issue_node_id = None
    cr_item_id = None
    if is_cr and project.github_project_node_id:
        try:
            cfg = cfg_for_project(current_user, project)
            issue = gh.create_issue(
                title=f"[Change Request] {note_title}",
                body=f"**Change Request**\n\n{body.raw_notes}",
                cfg=cfg,
            )
            cr_issue_url = issue["url"]
            cr_issue_number = issue["number"]
            cr_issue_node_id = issue["node_id"]
            cr_item_id = gh.add_to_project(issue["node_id"], cfg=cfg)
            gh.update_project_status(cr_item_id, "Backlog", cfg=cfg)
        except Exception as e:
            log.warning(f"Could not create CR GitHub issue: {e}")

    note = models.MeetingNote(
        creator_id=current_user.id,
        project_id=project_id,
        note_type="change_request" if is_cr else "note",
        raw_notes=body.raw_notes,
        wiki_url=body.wiki_url or None,
        title=note_title,
        status="In Progress" if is_cr else "Draft",
        brd_generation_phase=0 if is_cr else None,
        # Change requests get their own issue; regular notes share the project issue
        github_issue_url=cr_issue_url if is_cr else project.github_issue_url,
        github_issue_number=cr_issue_number if is_cr else project.github_issue_number,
        github_issue_node_id=cr_issue_node_id if is_cr else project.github_issue_node_id,
        github_project_item_id=cr_item_id if is_cr else project.github_project_item_id,
    )
    db.add(note)
    db.flush()

    entry = models.NoteEntry(note_id=note.id, content=body.raw_notes)
    db.add(entry)
    db.commit()
    db.refresh(note)

    if is_cr:
        background_tasks.add_task(_full_cr_pipeline, note.id)

    return _note_dict(note, db)


@projects_notes_router.post("/{project_id}/enhance-against-prd")
def enhance_against_prd(
    project_id: str,
    body: schemas.EnhanceAgainstPrdRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Enhance change-request text against the project's approved PRD content."""
    project = db.query(models.Project).filter(
        models.Project.id == project_id,
        models.Project.creator_id == current_user.id,
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Find the Sent to Planner PRD for this project
    prd = (
        db.query(models.PrdDocument)
        .join(models.MeetingNote, models.MeetingNote.id == models.PrdDocument.note_id)
        .filter(
            models.MeetingNote.project_id == project_id,
            models.PrdDocument.status == "Sent to Planner",
        )
        .first()
    )
    if not prd or not prd.prd_draft:
        raise HTTPException(status_code=404, detail="No approved PRD found for this project")

    try:
        enhanced = agent.enhance_against_prd(body.raw_text, prd.prd_draft)
    except Exception as e:
        log.error(f"Enhance against PRD failed: {e}")
        raise HTTPException(status_code=502, detail="AI enhancement failed")

    return {"enhanced_text": enhanced}


@projects_notes_router.get("/{project_id}/notes", response_model=schemas.MeetingNotesListResponse)
def list_notes_for_project(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    project = db.query(models.Project).filter(
        models.Project.id == project_id,
        models.Project.creator_id == current_user.id,
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    notes = (
        db.query(models.MeetingNote)
        .filter(
            models.MeetingNote.project_id == project_id,
            models.MeetingNote.creator_id == current_user.id,
        )
        .order_by(models.MeetingNote.created_at.desc())
        .all()
    )
    return {"notes": [_note_dict(n, db) for n in notes], "total": len(notes)}


# ── Routes ────────────────────────────────────────────────────────────────────

@router.post("", response_model=schemas.MeetingNoteResponse, status_code=201)
def create_note(
    body: schemas.NotesEnhanceRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Save the first note entry (standalone, no project). No BRD generation — that happens on Mark as Ready."""
    note = models.MeetingNote(
        creator_id=current_user.id,
        project_id=body.project_id or None,
        raw_notes=body.raw_notes,
        wiki_url=body.wiki_url or None,
        title=_title_from_raw(body.raw_notes),
        status="Draft",
    )
    db.add(note)
    db.flush()

    entry = models.NoteEntry(note_id=note.id, content=body.raw_notes)
    db.add(entry)
    db.commit()
    db.refresh(note)
    return note


@router.post("/{note_id}/entries", response_model=schemas.MeetingNoteResponse, status_code=201)
def add_entry(
    note_id: str,
    body: schemas.NoteEntryCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Add another note entry while still in Draft. No BRD generation yet."""
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if note.status != "Draft":
        raise HTTPException(status_code=400, detail="Can only add notes while in Draft status")

    entry = models.NoteEntry(note_id=note.id, content=body.content)
    db.add(entry)
    if body.wiki_url is not None:
        note.wiki_url = body.wiki_url if body.wiki_url.strip() else None
    db.commit()
    db.refresh(note)
    return note


@router.get("/{note_id}/entries", response_model=schemas.NoteEntriesListResponse)
def list_entries(
    note_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    entries = (
        db.query(models.NoteEntry)
        .filter(models.NoteEntry.note_id == note_id)
        .order_by(models.NoteEntry.created_at.asc())
        .all()
    )
    return {"entries": entries, "total": len(entries)}


@router.post("/{note_id}/attachments", response_model=schemas.NoteAttachmentResponse, status_code=201)
def upload_attachment(
    note_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    uploads_dir = Path("uploads") / note_id
    uploads_dir.mkdir(parents=True, exist_ok=True)

    safe_name = re.sub(r"[^\w.\-]", "_", file.filename or "file")
    file_path = uploads_dir / safe_name
    content = file.file.read()
    file_path.write_bytes(content)

    mime = file.content_type or ""
    filename_lower = (file.filename or "").lower()
    content_text: str | None = None

    if "text" in mime:
        content_text = content.decode("utf-8", errors="ignore")
    elif "pdf" in mime:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            content_text = "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            pass
    elif (
        "wordprocessingml" in mime        # .docx
        or "msword" in mime               # .doc
        or filename_lower.endswith(".docx")
        or filename_lower.endswith(".doc")
    ):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            content_text = "\n".join(paragraphs) or None
        except Exception:
            pass

    attachment = models.NoteAttachment(
        note_id=note_id,
        filename=file.filename or safe_name,
        mime_type=mime,
        file_path=str(file_path),
        content_text=content_text or None,
    )
    db.add(attachment)
    db.commit()
    db.refresh(attachment)
    return {
        "id": attachment.id,
        "note_id": attachment.note_id,
        "filename": attachment.filename,
        "mime_type": attachment.mime_type,
        "has_text": attachment.content_text is not None,
        "created_at": attachment.created_at,
    }


@router.post("/{note_id}/mark-ready", response_model=schemas.MeetingNoteResponse, status_code=202)
def mark_ready(
    note_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """
    Move note to In Progress and fire the full BRD pipeline.
    If the note belongs to a project, all GitHub activity goes on the project's issue.
    Otherwise creates a standalone issue (legacy behaviour).
    """
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if note.status != "Draft":
        raise HTTPException(status_code=400, detail=f"Note is already {note.status}")

    cfg = cfg_for_user(current_user)

    if note.project_id:
        # Project-linked note: use the project's own GitHub board
        project = db.query(models.Project).filter(
            models.Project.id == note.project_id
        ).first()
        if project and project.github_issue_number:
            # Use the project's dedicated board config
            project_cfg = cfg_for_project(current_user, project)
            if project.github_project_item_id:
                try:
                    gh.update_project_status(project.github_project_item_id, "In Progress", cfg=project_cfg)
                except Exception as e:
                    log.warning(f"Board status update failed: {e}")

            preview = note.raw_notes[:80] + ("…" if len(note.raw_notes) > 80 else "")
            try:
                gh.add_comment(
                    project.github_issue_number,
                    f"### BRD Generation Started\n\n"
                    f"Analysing notes and generating Business Requirements Document…\n\n"
                    f"> {preview}",
                    cfg=project_cfg,
                )
            except Exception as e:
                log.warning(f"GitHub start comment failed: {e}")

            note.github_issue_url = project.github_issue_url
            note.github_issue_number = project.github_issue_number
            note.github_issue_node_id = project.github_issue_node_id
            note.github_project_item_id = project.github_project_item_id
    else:
        # Legacy standalone note: create its own GitHub issue on shared board
        preview = note.raw_notes[:60] + ("…" if len(note.raw_notes) > 60 else "")
        try:
            issue = gh.create_issue(
                title=f"BRD: {preview}",
                body="Analysing meeting notes and generating BRD…\n\nThis will be updated automatically.",
                cfg=cfg,
            )
            item_id = gh.add_to_project(issue["node_id"], cfg=cfg)
            gh.update_project_status(item_id, "In Progress", cfg=cfg)
            note.github_issue_url = issue["url"]
            note.github_issue_number = issue["number"]
            note.github_issue_node_id = issue["node_id"]
            note.github_project_item_id = item_id
        except Exception as e:
            log.warning(f"Standalone GitHub issue creation failed: {e}")

    note.status = "In Progress"
    note.brd_generation_phase = 0
    note.github_last_checked_at = datetime.utcnow()
    db.commit()
    db.refresh(note)

    background_tasks.add_task(_full_brd_pipeline, note.id)
    return note


@router.post("/{note_id}/assign-reviewer", response_model=schemas.MeetingNoteResponse)
def assign_reviewer(
    note_id: str,
    body: schemas.AssignReviewerRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Assign reviewer once BRD is ready. Moves board to In Review and notifies reviewer."""
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if note.status not in ("Pending Review", "In Review", "Changes Requested"):
        raise HTTPException(status_code=400, detail="BRD is not ready for reviewer assignment")

    cfg = cfg_for_user(current_user)

    if note.github_project_item_id:
        try:
            gh.update_project_status(note.github_project_item_id, "In Review", cfg=cfg)
        except Exception as e:
            log.warning(f"Board status update failed: {e}")

    # Append to reviewers list (no duplicates)
    reviewers = list(note.reviewers or [])
    existing_usernames = {r["github_username"].lower() for r in reviewers}
    is_new = body.reviewer_github_username.lower() not in existing_usernames
    if is_new:
        reviewers.append({
            "github_username": body.reviewer_github_username,
            "name": body.reviewer_name or body.reviewer_github_username,
            "status": "Pending",
        })
        note.reviewers = reviewers
        flag_modified(note, "reviewers")

    note.reviewer_github_username = body.reviewer_github_username
    note.reviewer_name = body.reviewer_name
    note.status = "In Review"
    db.commit()
    db.refresh(note)

    if is_new:
        issue_number = _get_project_issue_number(note, db)
        if issue_number:
            try:
                gh.add_comment(
                    issue_number,
                    f"@{body.reviewer_github_username} — this BRD has been assigned to you for review.\n\n"
                    f"The full Business Requirements Document is in the comment above.\n\n"
                    f"> Reply **`APPROVED`** to approve, or leave detailed feedback below and the AI agent will update the BRD automatically.",
                    cfg=cfg,
                )
            except Exception as e:
                log.warning(f"Failed to post assignment comment: {e}")

    return note


@router.post("/{note_id}/feedback", response_model=schemas.MeetingNoteResponse, status_code=202)
def submit_feedback(
    note_id: str,
    body: schemas.BrdFeedbackRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if not note.brd_draft:
        raise HTTPException(status_code=400, detail="BRD not yet generated")

    note.brd_generation_phase = 0
    note.status = "Changes Requested"
    db.commit()
    db.refresh(note)
    background_tasks.add_task(_update_brd, note.id, body.feedback)
    return note


@router.get("", response_model=schemas.MeetingNotesListResponse)
def list_notes(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    notes = (
        db.query(models.MeetingNote)
        .filter(models.MeetingNote.creator_id == current_user.id)
        .order_by(models.MeetingNote.created_at.desc())
        .all()
    )
    return {"notes": notes, "total": len(notes)}


@router.get("/{note_id}", response_model=schemas.MeetingNoteResponse)
def get_note(
    note_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    return _note_dict(note, db)


@router.patch("/{note_id}", response_model=schemas.MeetingNoteResponse)
def update_note(
    note_id: str,
    body: schemas.NoteEntryCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Update raw_notes on a note (used to edit Change Requests before Closed)."""
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if note.note_type == "change_request" and note.status == "Closed":
        raise HTTPException(status_code=409, detail="Cannot edit a Closed Change Request")

    note.raw_notes = body.content
    note.title = _title_from_raw(body.content)
    db.commit()
    db.refresh(note)
    return _note_dict(note, db)


@router.get("/{note_id}/versions", response_model=schemas.BrdVersionListResponse)
def list_versions(
    note_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    versions = (
        db.query(models.BrdVersion)
        .filter(models.BrdVersion.note_id == note_id)
        .order_by(models.BrdVersion.version_number.desc())
        .all()
    )
    return {"versions": versions, "total": len(versions)}


@router.post("/{note_id}/regenerate-cr-summary", response_model=schemas.MeetingNoteResponse, status_code=202)
def regenerate_cr_summary(
    note_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Regenerate the AI summary for a Change Request (only while In Progress or In Review)."""
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if note.note_type != "change_request":
        raise HTTPException(status_code=400, detail="Only Change Requests can use this endpoint")
    if note.status == "Closed":
        raise HTTPException(status_code=400, detail="Cannot regenerate summary for a Closed Change Request")

    note.status = "In Progress"
    note.brd_generation_phase = 0
    db.commit()
    db.refresh(note)

    background_tasks.add_task(_full_cr_pipeline, note.id)
    return _note_dict(note, db)


@router.post("/{note_id}/send-cr-to-planner", response_model=schemas.MeetingNoteResponse, status_code=202)
def send_cr_to_planner(
    note_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    """Finalise a Change Request: generate planner document, upload to GitHub, set Closed."""
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    if note.note_type != "change_request":
        raise HTTPException(status_code=400, detail="Only Change Requests can be sent to Planner via this endpoint")
    if note.status not in ("In Review",):
        raise HTTPException(status_code=400, detail="Change Request must be In Review before sending to Planner")

    note.status = "In Progress"   # show processing state while planner doc is generated
    note.brd_generation_phase = 0
    db.commit()
    db.refresh(note)

    background_tasks.add_task(_cr_send_to_planner_task, note.id)
    return _note_dict(note, db)


@router.delete("/{note_id}", status_code=204)
def delete_note(
    note_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    note = db.query(models.MeetingNote).filter(
        models.MeetingNote.id == note_id,
        models.MeetingNote.creator_id == current_user.id,
    ).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    # Block deletion if the note has an approved PRD
    approved_prd = db.query(models.PrdDocument).filter(
        models.PrdDocument.note_id == note_id,
        models.PrdDocument.status.in_(["Approved", "Sent to Planner"]),
    ).first()
    if approved_prd:
        raise HTTPException(
            status_code=409,
            detail="Cannot delete note: it has an approved PRD.",
        )

    db.delete(note)
    db.commit()
